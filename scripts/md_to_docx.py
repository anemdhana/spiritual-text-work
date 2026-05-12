import sys
from pathlib import Path
from markdown import markdown
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
import re

def md_to_docx(input_md: str, output_docx: str = None):
    input_path = Path(input_md)
    if not input_path.exists():
        print(f"Error: File '{input_md}' not found!")
        return

    if output_docx is None:
        output_docx = input_path.with_suffix('.docx')

    # Read Markdown
    with open(input_path, encoding='utf-8') as f:
        md_content = f.read()

    # Convert Markdown to HTML (same extensions as your PDF version)
    html_body = markdown(md_content, extensions=['extra', 'tables', 'fenced_code'])

    # Parse HTML
    soup = BeautifulSoup(html_body, 'html.parser')

    # Create Document
    doc = Document()

    # ====================== PAGE SETUP ======================
    section = doc.sections[0]

    # Minimal margins (matching ~0.8cm in PDF)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    # Page size A4
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)

    # ====================== STYLES SETUP ======================
    # Base font and paragraph style (Georgia-like serif)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Georgia'
    font.size = Pt(11.5)
    font.color.rgb = RGBColor(44, 44, 44)  # #2c2c2c

    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    paragraph_format.line_spacing = 1.65
    paragraph_format.space_after = Pt(8)

    # Heading styles
    for i in range(1, 5):
        try:
            heading_style = doc.styles[f'Heading {i}']
        except KeyError:
            heading_style = doc.styles.add_style(f'Heading {i}', WD_STYLE_TYPE.PARAGRAPH)
        
        h_font = heading_style.font
        h_font.name = 'Georgia'
        h_font.color.rgb = RGBColor(44, 62, 80)  # #2c3e50
        h_font.bold = True
        if i == 1:
            h_font.size = Pt(22)   # approx 1.85em
            heading_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif i == 2:
            h_font.size = Pt(18)
        elif i == 3:
            h_font.size = Pt(15)
        else:
            h_font.size = Pt(13)

        heading_style.paragraph_format.space_before = Pt(20)
        heading_style.paragraph_format.space_after = Pt(10)

    # Blockquote style
    try:
        quote_style = doc.styles['Intense Quote']
    except KeyError:
        quote_style = doc.styles.add_style('Intense Quote', WD_STYLE_TYPE.PARAGRAPH)
    quote_font = quote_style.font
    quote_font.name = 'Georgia'
    quote_font.italic = True
    quote_font.color.rgb = RGBColor(85, 85, 85)  # #555
    quote_style.paragraph_format.left_indent = Inches(0.5)
    quote_style.paragraph_format.space_before = Pt(14)
    quote_style.paragraph_format.space_after = Pt(14)

    # Code style
    try:
        code_style = doc.styles.add_style('Code', WD_STYLE_TYPE.CHARACTER)
    except KeyError:
        code_style = doc.styles['Code']
    code_font = code_style.font
    code_font.name = 'Consolas'  # or 'Courier New'
    code_font.size = Pt(10.5)
    code_font.color.rgb = RGBColor(0, 0, 0)

    # ====================== CONTENT CONVERSION ======================
    for element in soup.find_all(recursive=False):  # top-level elements
        if element.name == 'h1':
            p = doc.add_heading(element.get_text(), level=1)
        elif element.name == 'h2':
            p = doc.add_heading(element.get_text(), level=2)
        elif element.name == 'h3':
            p = doc.add_heading(element.get_text(), level=3)
        elif element.name == 'h4':
            p = doc.add_heading(element.get_text(), level=4)
        elif element.name == 'p':
            p = doc.add_paragraph()
            _add_formatted_runs(p, element)
        elif element.name == 'blockquote':
            p = doc.add_paragraph(element.get_text().strip(), style='Intense Quote')
        elif element.name == 'pre':
            # Code block
            code_text = element.get_text()
            p = doc.add_paragraph(code_text, style='Normal')
            for run in p.runs:
                run.style = 'Code'
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(10)
        elif element.name == 'ul':
            for li in element.find_all('li', recursive=False):
                p = doc.add_paragraph(li.get_text().strip(), style='List Bullet')
                _add_formatted_runs(p, li)
        elif element.name == 'ol':
            for li in element.find_all('li', recursive=False):
                p = doc.add_paragraph(li.get_text().strip(), style='List Number')
                _add_formatted_runs(p, li)
        elif element.name == 'table':
            _add_table(doc, element)
        elif element.name == 'img':
            try:
                img_src = element.get('src')
                if img_src:
                    img_path = input_path.parent / img_src
                    if img_path.exists():
                        doc.add_picture(str(img_path), width=Inches(6.0))
                    else:
                        print(f"Warning: Image not found: {img_src}")
            except Exception as e:
                print(f"Warning: Could not add image: {e}")
        else:
            text = element.get_text().strip()
            if text:
                p = doc.add_paragraph(text)

    # ====================== PAGE NUMBERS (FOOTER) ======================
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = ""
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run("Page ")
    run.font.size = Pt(9.5)
    run.font.name = 'Georgia'
    run.font.color.rgb = RGBColor(102, 102, 102)

    run = footer_para.add_run("X of Y")
    run.font.size = Pt(9.5)
    run.font.name = 'Georgia'
    run.font.color.rgb = RGBColor(102, 102, 102)

    print(f"✅ DOCX created successfully:")
    print(f"   {output_docx}")
    print("   Note: Open in Word and update fields (Ctrl+A → F9) for accurate page numbers.")
    doc.save(output_docx)


def main():
    if len(sys.argv) < 2:
        print("Usage: python md_to_docx.py your_file.md [output.docx]")
        sys.exit(1)

    input_file = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else None
    md_to_docx(input_file, output_file)


def _add_formatted_runs(paragraph, element):
    """Add inline formatting (bold, italic, code, links, etc.)"""
    for child in element.children:
        if child.name is None:
            text = child.string or ''
            if text.strip():
                run = paragraph.add_run(text)
        elif child.name == 'strong' or child.name == 'b':
            run = paragraph.add_run(child.get_text())
            run.bold = True
        elif child.name == 'em' or child.name == 'i':
            run = paragraph.add_run(child.get_text())
            run.italic = True
        elif child.name == 'code':
            run = paragraph.add_run(child.get_text())
            run.style = 'Code'
        elif child.name == 'a':
            run = paragraph.add_run(child.get_text())
        else:
            run = paragraph.add_run(child.get_text())


def _add_table(doc, table_elem):
    """Convert HTML table to docx table"""
    rows = table_elem.find_all('tr')
    if not rows:
        return

    cols = len(rows[0].find_all(['td', 'th']))
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = 'Table Grid'

    for i, row in enumerate(rows):
        cells = row.find_all(['td', 'th'])
        for j, cell in enumerate(cells):
            docx_cell = table.cell(i, j)
            p = docx_cell.paragraphs[0]
            p.text = cell.get_text().strip()
            if cell.name == 'th':
                for run in p.runs:
                    run.bold = True

    table.autofit = True


if __name__ == "__main__":
    main()
